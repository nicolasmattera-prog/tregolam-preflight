import os, re, difflib
from docx import Document
from docx.shared import RGBColor
from concurrent.futures import ThreadPoolExecutor
from openai import OpenAI
from dotenv import load_dotenv

# INTEGRACIÓN CON EL VALIDADOR APROBADO
# LOCAL (ejecución directa / Streamlit / tests locales)
# from validador import corregir_bloque_con_seguridad
# from regex_rules import RULES

# RENDER / PRODUCCIÓN (Gunicorn, contenedor Linux)
from scripts.validador import corregir_bloque_con_seguridad
from scripts.regex_rules import RULES
# ===========================================================


# ---------- CONFIGURACIÓN ----------
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
MODEL_MINI = "gpt-4o-mini"

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
INPUT_FOLDER = os.path.join(BASE_DIR, "entrada")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "salida")
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# ---------- PROMPTS ULTRA-ESTRICTOS (ÍNTEGROS) ----------
PROMPT_F1 = """
Eres un CORRECTOR ORTOGRÁFICO Y TIPOGRÁFICO de texto ya existente.  
Tu única tarea es aplicar, SIN EXCEPCIONES, las reglas que se listan a continuación:

1. Diálogos: Raya de apertura (—) pegada al texto. Raya de inciso pegada al texto (—dijo Rubén). Puntuación siempre después de la raya de cierre: —dijo—.
2. Mayúsculas: Corrige capitalización sin tocar siglas ni acrónimos.
3. Ortografía: Tildes, diéresis, v/b, haches y concordancia simple.
4. Signos: Quita repeticiones (,, !!, ??).
5. VOCATIVO: Coma obligatoria (ej: «Marta, cierra la puerta»).
6. ESPACIOS APERTURA: UN espacio entre la palabra anterior y el signo de apertura (¿, ¡, «).
7. PEGOTES: UN espacio después de punto, coma, etc. Separa «autenticidad.Los» -> «autenticidad. Los».
8. GRAMÁTICA: Corrige "si + habría" por "si + hubiera/hubiese".

RESTRICCIONES: No cambies palabras correctas, no añadas frases, no pongas comentarios.
"""

PROMPT_F2 = """Eres editor literario. Mejora la agilidad verbal:
1. GERUNDIOS DE POSTERIORIDAD: 'terminó, generando' -> 'terminó y generó'.
2. VOZ PASIVA: Cámbiala a activa REORDENANDO la frase.
3. ESTRUCTURAS PESADAS: Mejora el flujo natural.
4. LIMPIEZA LINGÜÍSTICA: Corrige queísmo/dequeísmo y concordancia de colectivos.

REGLA DE ORO: Respeta escrupulosamente los espacios en cifras y comillas de la fase anterior.
RESTRICCIÓN ABSOLUTA:
- No añadas ninguna raya (—) que no exista en el texto original.
- Si el texto original no es un diálogo, no pongas rayas.
- Si incumples esto, el cambio se rechaza."""

# ---------- MOTOR DE REGEX (ÍNTEGRO) ----------

def aplicar_regex_editorial(texto):
    if not texto:
        return ""

    texto = texto.replace('\\xa0', ' ').replace('\\u202f', ' ')
    texto = texto.replace('\xa0', ' ').replace('\u202f', ' ')

    for patron, reemplazo, nombre_regla in RULES:
        try:
            texto = patron.sub(reemplazo, texto)
        except Exception as e:
            print(f"Error aplicando {nombre_regla}: {e}")

    return texto


# ---------- FUNCIONES TÉCNICAS (ÍNTEGRAS) ----------

def _tokenize(txt):
    return re.findall(r'(\S+)([ \t\u00A0\r\n]*)', txt, re.UNICODE)

def limpieza_residuos_chat(texto):
    patrones_basura = [r"^claro, aquí tienes.*?:", r"^aquí está el texto.*?:", r"^he corregido.*?:", r"espero que te sirva.*$"]
    for patron in patrones_basura:
        texto = re.sub(patron, "", texto, flags=re.IGNORECASE | re.MULTILINE)
    return texto.strip().strip('"')

def eliminar_inserciones_largas(original, corregido, max_palabras=3):
    orig_tok = _tokenize(original)
    corr_tok = _tokenize(corregido)
    s = difflib.SequenceMatcher(None, [w for w, _ in orig_tok], [w for w, _ in corr_tok])
    out = []
    for tag, i1, i2, j1, j2 in s.get_opcodes():
        if tag == 'insert' and (j2 - j1) > max_palabras: continue
        if tag == 'replace' and (j2 - j1) > (i2 - i1) + max_palabras:
            for i in range(i1, i2):
                pal, esp = orig_tok[i]
                out.append(pal + esp)
            continue
        for j in range(j1, j2):
            pal, esp = corr_tok[j]
            out.append(pal + esp)
    return ''.join(out)

# ---------- MAPPING Y GUARDADO (CON LÓGICA DE COLOR EXPLÍCITA) ----------

def aplicar_cambios_quirurgicos(parrafo, original, corregido, score):
    if original.strip() == corregido.strip(): return
    era_cursiva = any(run.italic for run in parrafo.runs)
    for run in parrafo.runs: run.text = ""

    orig_words = [w for w, _ in _tokenize(original)]
    corr_list = _tokenize(corregido)
    corr_words = [w for w, e in corr_list]

    s = difflib.SequenceMatcher(None, orig_words, corr_words)
    
    for tag, i1, i2, j1, j2 in s.get_opcodes():
        segmento = "".join([w + e for w, e in corr_list[j1:j2]])
        if not segmento: continue
        
        run = parrafo.add_run(segmento)
        run.font.name = 'Garamond'
        # Intentar mantener el tamaño original
        try:
            run.font.size = parrafo.runs[0].font.size if parrafo.runs else None
        except:
            pass
        run.italic = era_cursiva
        
        # LÓGICA DE COLORES MEJORADA (MÁS EXPLÍCITA)
        if tag in ('replace', 'insert'):
            if score < 90: 
                run.font.color.rgb = RGBColor(255, 165, 0)  # Naranja (Dudoso)
            else:  # score >= 90
                if tag == 'replace':
                    run.font.color.rgb = RGBColor(0, 0, 180)  # Azul (Seguro)
                else:  # insert
                    run.font.color.rgb = RGBColor(180, 0, 0)  # Rojo (Seguro)
        else:
            run.font.color.rgb = RGBColor(0, 0, 0) # Negro (Sin cambios)

def procesar_archivo(name):
    print(f"🚀 Iniciando corrección modular: {name}")
    doc = Document(os.path.join(INPUT_FOLDER, name))
    
    # Recopilar párrafos y tablas
    objetivos = [p for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if p.text.strip(): objetivos.append(p)

    textos_orig = [p.text for p in objetivos]
    
    # Procesamiento paralelo llamando al validador aprobado
    with ThreadPoolExecutor(max_workers=8) as exe:
        # Se envían los prompts y la función regex para asegurar consistencia
        resultados_tuplas = list(exe.map(
            lambda t: corregir_bloque_con_seguridad(t, PROMPT_F1, PROMPT_F2, aplicar_regex_editorial), 
            textos_orig
        ))

    for p, orig, (corr, score) in zip(objetivos, textos_orig, resultados_tuplas):
        aplicar_cambios_quirurgicos(p, orig, corr, score)

    doc.save(os.path.join(OUTPUT_FOLDER, name))
    print(f"✅ Archivo guardado con éxito: {name}")

if __name__ == "__main__":
    archivos = [f for f in os.listdir(INPUT_FOLDER) if f.endswith(".docx") and not f.startswith("~$")]
    if not archivos:
        print("No se encontraron archivos .docx en la carpeta 'entrada'.")
    else:
        for a in archivos: 
            try:
                procesar_archivo(a)
            except Exception as e:
                print(f"❌ Error en {a}: {e}")

