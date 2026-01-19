#!/usr/bin/env python3
import os, re, difflib
from docx import Document
from docx.shared import RGBColor
from concurrent.futures import ThreadPoolExecutor
from openai import OpenAI
from dotenv import load_dotenv
import time

from token_monitor import log_tokens

# ---------- CONFIG ----------
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=API_KEY)
MODEL_MINI = "gpt-4o-mini"
MODEL_FULL = "gpt-4o"

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
INPUT_FOLDER = os.path.join(BASE_DIR, "entrada")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "salida")
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# ---------- PROMPT ----------
PROMPT = """
Eres un CORRECTOR ORTOGRÁFICO Y TIPOGRÁFICO de texto ya existente.  
Tu única tarea es aplicar, SIN EXCEPCIONES, las reglas que se listan a continuación. Nada de lo que no se mencione está permitido.

1. Números: 4 cifras seguidas (4000). 5 o más cifras: espacio cada 3 (20 000). Años: juntos (2026). Porcentajes: espacio antes del % (20 %).
2. Unidades y símbolos: Espacio entre cantidad y símbolo (12 kg, 45 °C, 60 %). Sin plural en símbolos (kg, %, cm).
3. Abreviaturas: EE. UU., a. C., n.º, D.ª, Sr.
4. Diálogos y citas: Raya de apertura (—) pegada al texto. Raya de inciso pegada al texto (—dijo Rubén). Puntuación siempre después de la raya de cierre: —dijo—. / «eso».
5. Comillas: Sustituye CUALQUIER tipo de comilla doble (ya sean rectas " ", curvas de apertura “ o curvas de cierre ”) por comillas latinas « » siempre.
7. Mayúsculas: Corrige capitalización sin tocar siglas ni acrónimos.
8. Ortografía general: Tildes, diéresis, haches, v/b, y/ll, etc.
9. Signos de puntuación: Quita repeticiones (,, !!, ??).
10. VOCATIVO: Coma obligatoria para separar el vocativo (ej: «Marta, cierra la puerta», «Hoy, amigos, celebramos»).

11. REGLA DE ESPACIOS DE APERTURA (OBLIGATORIA): 
    - Siempre debe haber UN espacio entre la palabra anterior y el signo de apertura.
    - Ejemplo correcto: «palabra ¿», «palabra ¡», «palabra «».
    - NUNCA pegues el signo de apertura a la palabra que le precede.

12. REGLA DE ESPACIOS DE CIERRE Y PEGOTES:
    - Siempre debe haber UN espacio después de punto, coma, punto y coma y dos puntos.
    - Si dos frases están pegadas por un punto (ej: «autenticidad.Los»), separa OBLIGATORIAMENTE con un espacio: «autenticidad. Los».
    - Nunca pegues una palabra inmediatamente después de un signo de puntuación de cierre.

RESTRICCIONES ABSOLUTAS (infringir cualquera anula la corrección):
- No cambies ni una palabra que esté bien escrita.
- No añadas, suprimas ni reordenes frases.
- No introduzcas comentarios, explicaciones ni ejemplos.
- No uses asteriscos ni otros marcadores.
- No generes párrafos nuevos ni líneas en blanco extra.
- No corrijas estilo, solo errores ortográficos/tipográficos.
- Mantén la longitud del texto lo más cercana posible al original.
- Cumple la regla 12 al pie: nunca quites el espacio tras . , ; : y nunca pegues palabras a esos signos.
"""

# ---------- PROMPT MICRO VERBAL (FASE 2) ----------
PROMPT_MICRO_VERBAL = """
Eres un corrector gramatical especializado en uso verbal en español.

Si el fragmento contiene un uso verbal incorrecto
(infinitivo exhortativo, gerundio mal empleado,
participio incorrecto o construcción verbal impropia),
corrige ÚNICAMENTE el verbo o la construcción verbal mínima necesaria.

NO resumas.
NO reescribas el párrafo completo.
NO añadas ni elimines información.
NO cambies el significado.
Si no hay error verbal, devuelve el fragmento EXACTAMENTE igual.

Devuelve solo el fragmento corregido.
"""

# ---------- MICRO VERBAL (FASE 2) ----------
def corregir_verbal_micro(fragmento):
    if not fragmento.strip():
        return fragmento
    try:
        res = client.chat.completions.create(
            model=MODEL_MINI,
            messages=[
                {"role": "system", "content": PROMPT_MICRO_VERBAL},
                {"role": "user", "content": fragmento}
            ],
            temperature=0
        )
        log_tokens(model=MODEL_MINI, usage=res.usage, tag="micro_verbal")
        r = res.choices[0].message.content.strip()
        return r if r else fragmento
    except Exception as e:
        print("⚠️ micro_verbal error:", e)
        return fragmento
    
# ---------- LIMPIEZA ----------
def limpieza_mecanica(texto):
    if not texto:
        return ""
    texto = re.sub(r'([.,;:?!»])([a-zA-ZáéíóúÁÉÍÓÚ0-9])', r'\1 \2', texto)
    texto = re.sub(r'([a-zA-ZáéíóúÁÉÍÓÚ0-9])([¿¡«])', r'\1 \2', texto)
    return re.sub(r' +', ' ', texto).strip()

# ---------- NUEVA CAPA PYTHON ----------
def correcciones_gramaticales_seguras(texto):
    reglas = [
        (r'\bsi habría\b', 'si hubiera'),
        (r'\bhabían\s+([a-záéíóúñ]+)', r'había \1'),
        (r'\bhubieron\s+([a-záéíóúñ]+)', r'hubo \1'),
        (r'\bpuede se\b', 'puede ser'),
        (r'\bparece estar mal redactar\b', 'parece estar mal redactado'),
        (r'\binsistió en de\b', 'insistió en'),
        (r'\bdepende que\b', 'depende de que'),
    ]
    for patron, reemplazo in reglas:
        texto = re.sub(patron, reemplazo, texto, flags=re.IGNORECASE)
    return texto

# ---------- CORRECCIÓN CON REINTENTOS ----------
def corregir_bloque(texto):
    if not texto.strip():
        return texto
    
    intentos_maximos = 3
    for intento in range(intentos_maximos):
        try:
            # ---- FASE 1: PROMPT BASE ----
            res = client.chat.completions.create(
                model=MODEL_MINI,
                messages=[
                    {"role": "system", "content": PROMPT},
                    {"role": "user", "content": texto}
                ],
                temperature=0,
                timeout=40
            )
            log_tokens(model=MODEL_MINI, usage=res.usage, tag="mini")
            r = res.choices[0].message.content.strip()

            # ---- FALLBACK A MODELO FULL (Si la respuesta es muy corta o falla) ----
            if not r or len(r) < len(texto) * 0.80:
                res = client.chat.completions.create(
                    model=MODEL_FULL,
                    messages=[
                        {"role": "system", "content": PROMPT},
                        {"role": "user", "content": texto}
                    ],
                    temperature=0
                )
                log_tokens(model=MODEL_FULL, usage=res.usage, tag="fallback_full")
                r = res.choices[0].message.content.strip()

            # ---- FASE 2: MICRO VERBAL ----
            # r2 = corregir_verbal_micro(r)
            # if r2 and r2.strip():
               #  r = r2

            # ---- NIVEL PYTHON SEGURO ----
            r = correcciones_gramaticales_seguras(r)

            # ---- LIMPIEZA FINAL ----
            return limpieza_mecanica(r)

        except Exception as e:
            error_msg = str(e).lower()
            if "rate_limit" in error_msg or "429" in error_msg:
                tiempo_espera = (intento + 1) * 3 # Esperamos 3, 6, 9 segundos
                print(f"⏳ Límite OpenAI alcanzado. Reintentando en {tiempo_espera}s...")
                time.sleep(tiempo_espera)
                continue
            
            print(f"⚠️ Error inesperado: {e}")
            return texto # En error crítico, devolvemos original
            
    return texto

# ---------- PINTADO ----------
def pintar_quirurgico(parrafo, original, corregido):
    if original == corregido:
        return
    era_cursiva = any(run.italic for run in parrafo.runs)
    for run in parrafo.runs:
        run.text = ""

    s = difflib.SequenceMatcher(None, original.split(), corregido.split())
    for tag, i1, i2, j1, j2 in s.get_opcodes():
        palabras = corregido.split()[j1:j2]
        if not palabras:
            continue
        txt = " ".join(palabras) + " "
        run = parrafo.add_run(txt)
        run.font.name = 'Garamond'
        run.italic = era_cursiva
        run.font.color.rgb = RGBColor(0, 0, 180) if tag in ('replace', 'insert') else RGBColor(0, 0, 0)

# ---------- PROCESADO ----------
def procesar_archivo(name):
    print(f"📄 Procesando: {name}")
    doc = Document(os.path.join(INPUT_FOLDER, name))
    parrafos = [p for p in doc.paragraphs]
    textos_originales = [p.text for p in parrafos]

    with ThreadPoolExecutor(max_workers=3) as exe:
        resultados = list(exe.map(corregir_bloque, textos_originales))

    for p, corregido in zip(parrafos, resultados):
        pintar_quirurgico(p, p.text, corregido)

    doc.save(os.path.join(OUTPUT_FOLDER, name.replace(".docx", "_CORREGIDO.docx")))
    print("✔ Finalizado.")

# ---------- MODULO DE COMPROBACION (OPTIMIZADO) ----------
def comprobar_archivo(name):
    print(f"🔍 Iniciando comprobación rápida de: {name}")
    ruta_entrada = os.path.join(INPUT_FOLDER, name)
    doc = Document(ruta_entrada)
    
    # Extraemos solo párrafos con texto
    p_objetos = [p for p in doc.paragraphs if p.text.strip()]
    textos_originales = [p.text.strip() for p in p_objetos]

    print(f"📡 Analizando {len(textos_originales)} párrafos en paralelo...")
    
    # PROCESAMIENTO EN PARALELO (Crucial para que no sea lento)
    with ThreadPoolExecutor(max_workers=3) as exe:
        resultados = list(exe.map(corregir_bloque, textos_originales))

    informe = [f"AUDITORÍA DE CALIDAD: {name}\n" + "=" * 40 + "\n"]
    encontrados = 0

    for i, (ori, limpio) in enumerate(zip(textos_originales, resultados)):
        ori_n = normalizar_para_auditoria(ori)
        lim_n = normalizar_para_auditoria(limpio)

        if ori_n != lim_n:
            encontrados += 1
            informe.append(f"📍 PÁRRAFO {i+1}")
            informe.append(f"ORIGINAL:   {ori}")
            informe.append(f"SUGERENCIA: {limpio}")
            informe.append("-" * 20)

    nombre_txt = f"VALIDACION_{name.replace('.docx', '')}.txt"
    ruta_txt = os.path.join(OUTPUT_FOLDER, nombre_txt)
    
    with open(ruta_txt, "w", encoding="utf-8") as f:
        f.write("\n".join(informe))
    
    print(f"✅ Informe generado: {nombre_txt} con {encontrados} avisos.")
    return nombre_txt

# ---------- MAIN ----------
if __name__ == "__main__":
    archivos = [f for f in os.listdir(INPUT_FOLDER) if f.endswith(".docx")]
    for a in archivos:
        procesar_archivo(a)




